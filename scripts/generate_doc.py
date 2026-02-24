#!/usr/bin/env python3
"""Generate Voice Navigator Technical Setup Word Document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0x1B, 0x5E, 0x8C)
    h.font.name = 'Calibri'

# Helper to add a styled table
def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()

def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run('Note: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.italic = True

# ── Title Page ──────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Voice Navigator for Salesforce')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x8C)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Technical Components & Installation Guide')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('Version 1.1.0  |  Released  |  API Version 59.0')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()
meta2 = doc.add_paragraph()
meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta2.add_run('Package Install URL (Production):\nhttps://login.salesforce.com/packaging/installPackage.apexp?p0=04tJ1000000t3WbIAI')
run.font.size = Pt(10)
run.font.name = 'Consolas'

meta3 = doc.add_paragraph()
meta3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta3.add_run('Package Install URL (Sandbox):\nhttps://test.salesforce.com/packaging/installPackage.apexp?p0=04tJ1000000t3WbIAI')
run.font.size = Pt(10)
run.font.name = 'Consolas'

doc.add_page_break()

# ── Table of Contents ───────────────────────────────────────────────────
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Overview',
    '2. Architecture',
    '3. Technical Components',
    '4. Installation Methods',
    '5. Post-Installation: Add to Utility Bar (Recommended)',
    '6. Post-Installation: Add to Home Page',
    '7. Post-Installation: Add to Record Page',
    '8. Microphone & Browser Setup',
    '9. Voice Commands Reference',
    '10. Security & Permissions',
    '11. Troubleshooting',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ── 1. Overview ─────────────────────────────────────────────────────────
doc.add_heading('1. Overview', level=1)
doc.add_paragraph(
    'Voice Navigator is a voice-controlled navigation component for Salesforce Lightning Experience. '
    'It allows users to navigate Salesforce by speaking natural commands such as "Go to accounts", '
    '"Create new contact", "Find Acme", or "Open users setup page".'
)
doc.add_paragraph(
    'The component uses the browser\'s Web Speech API for speech recognition and works entirely '
    'within Salesforce — no external APIs, no API keys, and no additional configuration required.'
)

doc.add_heading('Key Facts', level=2)
add_table(
    ['Property', 'Value'],
    [
        ['Platform', 'Salesforce Lightning Experience'],
        ['API Version', '59.0'],
        ['Package Type', 'Unlocked Package (no namespace)'],
        ['Package Version ID', '04tJ1000000t3WbIAI'],
        ['Version', '1.1.0 (Released)'],
        ['Status', 'Production-ready'],
        ['Code Coverage', '93%'],
        ['External Dependencies', 'None'],
        ['API Keys Required', 'None'],
        ['Browser Required', 'Google Chrome (desktop)'],
        ['Language', 'English (US)'],
        ['Navigation Behavior', 'Opens pages in new tab (utility bar stays open)'],
    ]
)

# ── 2. Architecture ────────────────────────────────────────────────────
doc.add_heading('2. Architecture', level=1)
doc.add_paragraph(
    'Voice Navigator consists of three main layers that work together:'
)

doc.add_heading('Architecture Diagram', level=2)
add_code_block(
    '+------------------------------------------------------------------+\n'
    '|                Salesforce Lightning Experience                    |\n'
    '|                                                                  |\n'
    '|  +----------------------------+  +----------------------------+  |\n'
    '|  |  LWC: voiceNavigator       |  |  Apex: VoiceNavigator-    |  |\n'
    '|  |                            |  |  Controller                |  |\n'
    '|  |  - Command parsing engine  |<>|                            |  |\n'
    '|  |  - Navigation (new tab)    |  |  - getObjectApiNames()    |  |\n'
    '|  |  - Fuzzy matching          |  |  - getObjectFields()      |  |\n'
    '|  |  - Search UI               |  |  - searchRecords()        |  |\n'
    '|  +------------|---------------+  +----------------------------+  |\n'
    '|               | postMessage             | SOSL / Schema API     |\n'
    '|               v                         v                       |\n'
    '|  +----------------------------+  +----------------------------+  |\n'
    '|  |  Static Resource:          |  |  Salesforce Database       |  |\n'
    '|  |  VoiceRecognitionSR        |  |  - Standard Objects        |  |\n'
    '|  |  (Hidden iframe,           |  |  - Custom Objects          |  |\n'
    '|  |   same-origin, no CSP)     |  |  - Record Search           |  |\n'
    '|  |  - Web Speech API          |  |                            |  |\n'
    '|  |  - SpeechSynthesis API     |  |                            |  |\n'
    '|  +----------------------------+  +----------------------------+  |\n'
    '+------------------------------------------------------------------+'
)

doc.add_heading('Data Flow', level=2)
doc.add_paragraph(
    '1. User clicks microphone button (or presses Option+Space / Ctrl+Space)\n'
    '2. LWC sends "start" message to hidden Static Resource iframe via postMessage\n'
    '3. Static resource activates Web Speech API and captures microphone audio\n'
    '4. Browser processes speech and returns transcript + confidence score\n'
    '5. Static resource sends transcript back to LWC via postMessage\n'
    '6. LWC command parser matches the transcript against known commands\n'
    '7. Navigation opens target page in a NEW TAB (utility bar stays open)\n'
    '8. Voice feedback speaks confirmation via SpeechSynthesis API'
)

doc.add_heading('Why a Static Resource (not Visualforce)?', level=2)
doc.add_paragraph(
    'Salesforce\'s Locker Service (security sandbox for LWC) blocks direct access to the browser\'s '
    'Web Speech API. The speech recognition code runs in a hidden iframe loaded from a Static Resource. '
    'Unlike Visualforce pages, static resources are served from the same origin as Lightning, which '
    'eliminates Content Security Policy (CSP) framing errors that occur in sandboxes. '
    'No security settings need to be changed — it works out of the box in all orgs.'
)

# ── 3. Technical Components ─────────────────────────────────────────────
doc.add_heading('3. Technical Components', level=1)

doc.add_heading('3.1 Lightning Web Component: voiceNavigator', level=2)
add_table(
    ['File', 'Size', 'Purpose'],
    [
        ['voiceNavigator.js', '39 KB', 'Core logic: command parsing, navigation, fuzzy matching, search integration'],
        ['voiceNavigator.html', '7 KB', 'UI template: microphone button, status display, search results, help section'],
        ['voiceNavigator.css', '1.6 KB', 'Styling: button animations, pulse effect, status colors, layout'],
        ['voiceNavigator.js-meta.xml', '420 B', 'Metadata: exposed to App Page, Home Page, Utility Bar, Record Page'],
    ]
)

doc.add_paragraph('Placement Targets (where the component can be added):')
bullets = [
    'lightning__UtilityBar — Utility bar (recommended)',
    'lightning__HomePage — Home page',
    'lightning__AppPage — Lightning App pages',
    'lightning__RecordPage — Record detail pages',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Key capabilities built into the LWC:')
features = [
    '24 standard object mappings (Account, Contact, Opportunity, Lead, Case, etc.)',
    '120+ Setup page URL mappings (Users, Profiles, Flows, Apex Classes, etc.)',
    '80+ Object Manager subpage mappings (Fields, Page Layouts, Validation Rules, etc.)',
    'Levenshtein distance fuzzy matching for typo/accent tolerance',
    'Dynamic custom object discovery via Apex wire service',
    'Keyboard shortcut: Option+Space (Mac) / Ctrl+Space (Windows)',
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('3.2 Apex Controller: VoiceNavigatorController', level=2)
add_table(
    ['File', 'Size', 'Purpose'],
    [
        ['VoiceNavigatorController.cls', '5.6 KB', 'Server-side logic: object discovery, field metadata, record search'],
        ['VoiceNavigatorController.cls-meta.xml', '150 B', 'Metadata: API version 59.0, Active status'],
    ]
)

doc.add_paragraph('Apex Methods:')
add_table(
    ['Method', 'Cacheable', 'Purpose'],
    [
        ['getObjectApiNames()', 'Yes', 'Returns all custom objects in the org (label to API name mapping). Automatically discovers custom objects.'],
        ['getObjectFields(objectApiName)', 'Yes', 'Returns field metadata for any object. Respects Field-Level Security (FLS).'],
        ['searchRecords(objectApiName, searchTerm)', 'No', 'SOSL search for records by name. Supports object-scoped and cross-object search. Input sanitized.'],
    ]
)

doc.add_heading('3.3 Apex Test Class: VoiceNavigatorControllerTest', level=2)
add_table(
    ['File', 'Size', 'Purpose'],
    [
        ['VoiceNavigatorControllerTest.cls', '6 KB', '13 test methods covering all Apex methods (93% code coverage)'],
        ['VoiceNavigatorControllerTest.cls-meta.xml', '150 B', 'Metadata'],
    ]
)

doc.add_heading('3.4 Static Resource: VoiceRecognitionSR (Speech Engine)', level=2)
add_table(
    ['File', 'Size', 'Purpose'],
    [
        ['VoiceRecognitionSR.html', '3.2 KB', 'Speech recognition engine — Web Speech API + SpeechSynthesis in hidden iframe'],
        ['VoiceRecognitionSR.resource-meta.xml', '200 B', 'Metadata (contentType: text/html, cacheControl: Public)'],
    ]
)
doc.add_paragraph(
    'This static resource replaced the original Visualforce page (VoiceRecognitionPage) to eliminate '
    'CSP framing errors in sandboxes. The VF page is still included in the package for backward '
    'compatibility but is no longer used.'
)

doc.add_heading('3.5 Visualforce Page: VoiceRecognitionPage (Legacy)', level=2)
add_table(
    ['File', 'Size', 'Purpose'],
    [
        ['VoiceRecognitionPage.page', '3.2 KB', 'Legacy speech recognition bridge — replaced by VoiceRecognitionSR'],
        ['VoiceRecognitionPage.page-meta.xml', '250 B', 'Metadata'],
    ]
)

doc.add_paragraph('Communication Protocol (LWC and Static Resource iframe):')
add_table(
    ['Direction', 'Message', 'Purpose'],
    [
        ['LWC to VF', '{ action: "start" }', 'Begin listening for voice input'],
        ['LWC to VF', '{ action: "stop" }', 'Stop listening'],
        ['LWC to VF', '{ action: "speak", text: "..." }', 'Voice feedback via SpeechSynthesis'],
        ['VF to LWC', '{ type: "voiceResult", transcript, confidence }', 'Speech recognition result'],
        ['VF to LWC', '{ type: "voiceStatus", status }', 'Listening/ended status updates'],
        ['VF to LWC', '{ type: "voiceError", error }', 'Error events'],
    ]
)

doc.add_heading('3.6 Complete File Structure', level=2)
add_code_block(
    'voice-navigator/\n'
    '├── sfdx-project.json                           # Project config + package aliases\n'
    '├── manifest/\n'
    '│   └── package.xml                             # Metadata API deployment manifest\n'
    '├── scripts/\n'
    '│   ├── deploy.sh                               # Source deployment script\n'
    '│   ├── create-package.sh                       # Package creation script\n'
    '│   └── install-package.sh                      # Package install script\n'
    '└── force-app/main/default/\n'
    '    ├── classes/\n'
    '    │   ├── VoiceNavigatorController.cls         # Apex controller\n'
    '    │   ├── VoiceNavigatorController.cls-meta.xml\n'
    '    │   ├── VoiceNavigatorControllerTest.cls     # Test class (13 tests)\n'
    '    │   └── VoiceNavigatorControllerTest.cls-meta.xml\n'
    '    ├── lwc/voiceNavigator/\n'
    '    │   ├── voiceNavigator.html                  # UI template\n'
    '    │   ├── voiceNavigator.js                    # Core logic (39 KB)\n'
    '    │   ├── voiceNavigator.css                   # Styles\n'
    '    │   └── voiceNavigator.js-meta.xml           # Component metadata\n'
    '    ├── staticresources/\n'
    '    │   ├── VoiceRecognitionSR.html              # Speech engine (static resource)\n'
    '    │   └── VoiceRecognitionSR.resource-meta.xml\n'
    '    └── pages/\n'
    '        ├── VoiceRecognitionPage.page            # Legacy speech bridge (unused)\n'
    '        └── VoiceRecognitionPage.page-meta.xml'
)

doc.add_page_break()

# ── 4. Installation Methods ─────────────────────────────────────────────
doc.add_heading('4. Installation Methods', level=1)

doc.add_heading('Method 1: Unlocked Package Install (Recommended)', level=2)
doc.add_paragraph('Best for: Any org, non-technical admins, one-click install.')

doc.add_paragraph('For Production / Developer Edition orgs:')
doc.add_paragraph('Open this URL in your browser and follow the prompts:')
add_code_block('https://login.salesforce.com/packaging/installPackage.apexp?p0=04tJ1000000t3WbIAI')

doc.add_paragraph('For Sandbox orgs:')
add_code_block('https://test.salesforce.com/packaging/installPackage.apexp?p0=04tJ1000000t3WbIAI')

doc.add_paragraph('Or use the Salesforce CLI:')
add_code_block('sf package install --package 04tJ1000000t3WbIAI --target-org <org-alias> --wait 10')

doc.add_paragraph('Installation options when prompted:')
opts = [
    'Install for Admins Only — only System Administrators can use it',
    'Install for All Users — everyone in the org can use it (recommended)',
    'Install for Specific Profiles — choose which profiles get access',
]
for o in opts:
    doc.add_paragraph(o, style='List Bullet')

doc.add_heading('Method 2: SFDX Source Deploy (For Developers)', level=2)
doc.add_paragraph('Best for: Developers, CI/CD pipelines, customization before deployment.')
add_code_block(
    '# Clone the repository\n'
    'git clone <repository-url>\n'
    'cd voice-navigator\n\n'
    '# Authorize target org\n'
    'sf org login web --alias myorg\n\n'
    '# Deploy all components\n'
    './scripts/deploy.sh myorg\n\n'
    '# Or manually:\n'
    'sf project deploy start --target-org myorg --source-dir force-app --wait 10'
)

doc.add_heading('Method 3: Change Set (No CLI Required)', level=2)
doc.add_paragraph('Best for: Admins who prefer the Salesforce UI, sandbox-to-production deployment.')
doc.add_paragraph(
    'In the source org, create an Outbound Change Set named "Voice Navigator" and add these components:'
)
add_table(
    ['Component Type', 'Component Name'],
    [
        ['Apex Class', 'VoiceNavigatorController'],
        ['Apex Class', 'VoiceNavigatorControllerTest'],
        ['Lightning Web Component Bundle', 'voiceNavigator'],
        ['Static Resource', 'VoiceRecognitionSR'],
        ['Visualforce Page', 'VoiceRecognitionPage'],
    ]
)
doc.add_paragraph('Upload to the target org, then deploy from Inbound Change Sets with test class: VoiceNavigatorControllerTest.')

doc.add_page_break()

# ── 5. Utility Bar ──────────────────────────────────────────────────────
doc.add_heading('5. Post-Installation: Add to Utility Bar (Recommended)', level=1)
doc.add_paragraph(
    'The Utility Bar is the recommended placement. It makes Voice Navigator accessible as a pop-up '
    'panel from every page in the app, without taking up page layout space.'
)

doc.add_heading('Step-by-Step Instructions', level=2)
steps = [
    ('Step 1:', 'Go to Setup (click the gear icon in the top-right corner, then click "Setup").'),
    ('Step 2:', 'In the Quick Find box, search for "App Manager" and click on it.'),
    ('Step 3:', 'Find the Lightning App you want to add Voice Navigator to (e.g., "Sales", "Service", or your custom app).'),
    ('Step 4:', 'Click the dropdown arrow (▼) on the right side of that app row, then click "Edit".'),
    ('Step 5:', 'In the left sidebar of the App Settings, click "Utility Items (Desktop Only)".'),
    ('Step 6:', 'Click the "Add Utility Item" button.'),
    ('Step 7:', 'In the search box, type "voiceNavigator" and select it from the list.'),
    ('Step 8:', 'Configure the utility item with these settings:'),
]
for label, text in steps:
    p = doc.add_paragraph()
    run = p.add_run(label + ' ')
    run.bold = True
    p.add_run(text)

add_table(
    ['Setting', 'Value'],
    [
        ['Label', 'Voice Navigator'],
        ['Icon', 'voice (or microphone)'],
        ['Panel Width', '400'],
        ['Panel Height', '500'],
        ['Start Automatically', 'Unchecked (recommended)'],
    ]
)

p = doc.add_paragraph()
run = p.add_run('Step 9: ')
run.bold = True
p.add_run('Click "Save" to apply the changes.')

p = doc.add_paragraph()
run = p.add_run('Step 10: ')
run.bold = True
p.add_run('Refresh your browser. You should now see a "Voice Navigator" icon in the utility bar at the bottom of the screen.')

add_note('The utility bar is only visible in Lightning Experience on desktop browsers. It does not appear in the Salesforce mobile app.')

doc.add_page_break()

# ── 6. Home Page ────────────────────────────────────────────────────────
doc.add_heading('6. Post-Installation: Add to Home Page', level=1)
doc.add_paragraph(
    'Adding Voice Navigator to the Home page makes it visible when users first log in. '
    'It will appear as a component on the page layout.'
)

doc.add_heading('Step-by-Step Instructions', level=2)
steps = [
    ('Step 1:', 'Navigate to the Salesforce Home page.'),
    ('Step 2:', 'Click the gear icon in the top-right corner, then click "Edit Page". This opens the Lightning App Builder.'),
    ('Step 3:', 'In the left panel under "Components", search for "voiceNavigator".'),
    ('Step 4:', 'Drag the "voiceNavigator" component from the left panel onto your desired location on the page layout. The right sidebar is recommended so it doesn\'t interfere with the main content area.'),
    ('Step 5:', 'Click "Save" in the top-right corner.'),
    ('Step 6:', 'A prompt will appear asking about page activation. Click "Activate".'),
    ('Step 7:', 'Select "Assign as Org Default" to make it visible to all users, or choose specific app/profile assignments.'),
    ('Step 8:', 'Click "Save" to confirm activation.'),
]
for label, text in steps:
    p = doc.add_paragraph()
    run = p.add_run(label + ' ')
    run.bold = True
    p.add_run(text)

add_note('The component will only appear on the Home page. To have it on every page, use the Utility Bar method instead (Section 5).')

doc.add_page_break()

# ── 7. Record Page ──────────────────────────────────────────────────────
doc.add_heading('7. Post-Installation: Add to Record Page', level=1)
doc.add_paragraph(
    'You can add Voice Navigator to specific record pages (e.g., Account, Contact, Opportunity) '
    'so it\'s available when viewing or editing records.'
)

doc.add_heading('Step-by-Step Instructions', level=2)
steps = [
    ('Step 1:', 'Navigate to any record page (e.g., open an Account record).'),
    ('Step 2:', 'Click the gear icon in the top-right corner, then click "Edit Page". This opens the Lightning App Builder for that record page.'),
    ('Step 3:', 'In the left panel under "Components", search for "voiceNavigator".'),
    ('Step 4:', 'Drag the "voiceNavigator" component onto the page. The right sidebar or a narrow column works best.'),
    ('Step 5:', 'Click "Save".'),
    ('Step 6:', 'Click "Activate" when prompted.'),
    ('Step 7:', 'Choose the activation scope:'),
]
for label, text in steps:
    p = doc.add_paragraph()
    run = p.add_run(label + ' ')
    run.bold = True
    p.add_run(text)

activation_opts = [
    'Assign as Org Default — applies to all users viewing this object\'s records',
    'Assign as App Default — only applies within a specific app',
    'Assign as App, Record Type, and Profile — granular control',
]
for o in activation_opts:
    doc.add_paragraph(o, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('Step 8: ')
run.bold = True
p.add_run('Click "Save" to confirm.')

add_note('You need to repeat this for each object\'s record page where you want Voice Navigator. For universal access, the Utility Bar method (Section 5) is recommended.')

doc.add_page_break()

# ── 8. Browser Setup ───────────────────────────────────────────────────
doc.add_heading('8. Microphone & Browser Setup', level=1)

doc.add_heading('Granting Microphone Permission', level=2)
steps = [
    ('Step 1:', 'Click the microphone button in the Voice Navigator component.'),
    ('Step 2:', 'Chrome will show a popup asking "wants to use your microphone". Click "Allow".'),
    ('Step 3:', 'The component status will change to "Listening" and the button will pulse red.'),
]
for label, text in steps:
    p = doc.add_paragraph()
    run = p.add_run(label + ' ')
    run.bold = True
    p.add_run(text)

doc.add_heading('If You Accidentally Blocked the Microphone', level=2)
steps = [
    ('Step 1:', 'Click the lock/info icon in Chrome\'s address bar (left of the URL).'),
    ('Step 2:', 'Find "Microphone" in the permissions list.'),
    ('Step 3:', 'Change it from "Block" to "Allow".'),
    ('Step 4:', 'Refresh the page (F5 or Ctrl+R).'),
]
for label, text in steps:
    p = doc.add_paragraph()
    run = p.add_run(label + ' ')
    run.bold = True
    p.add_run(text)

doc.add_heading('Supported Browsers', level=2)
add_table(
    ['Browser', 'Speech Recognition', 'Voice Feedback', 'Status'],
    [
        ['Google Chrome (desktop)', 'Full support', 'Full support', 'Recommended'],
        ['Microsoft Edge (Chromium)', 'Full support', 'Full support', 'Supported'],
        ['Firefox', 'Not supported', 'Partial', 'Not recommended'],
        ['Safari', 'Limited', 'Full support', 'Not recommended'],
    ]
)

# ── 9. Voice Commands ──────────────────────────────────────────────────
doc.add_heading('9. Voice Commands Reference', level=1)

doc.add_heading('Object Navigation', level=2)
add_table(
    ['Say This', 'What Happens'],
    [
        ['"Go to accounts"', 'Opens the Account list view'],
        ['"Open contacts"', 'Opens the Contact list view'],
        ['"Navigate to opportunities"', 'Opens the Opportunity list view'],
        ['"Open leads"', 'Opens the Lead list view'],
        ['"Go to cases"', 'Opens the Case list view'],
        ['"Open campaigns"', 'Opens the Campaign list view'],
    ]
)

doc.add_heading('Record Creation', level=2)
add_table(
    ['Say This', 'What Happens'],
    [
        ['"Create new account"', 'Opens a new Account form'],
        ['"New contact"', 'Opens a new Contact form'],
        ['"Create opportunity"', 'Opens a new Opportunity form'],
        ['"Add a lead"', 'Opens a new Lead form'],
    ]
)

doc.add_heading('Record Search', level=2)
add_table(
    ['Say This', 'What Happens'],
    [
        ['"Find account Acme"', 'Searches Account records for "Acme"'],
        ['"Search contact John"', 'Searches Contact records for "John"'],
        ['"Find Acme"', 'Cross-object search (Account, Contact, Opportunity, Lead, Case)'],
    ]
)

doc.add_heading('Setup Navigation (120+ pages)', level=2)
add_table(
    ['Say This', 'What Happens'],
    [
        ['"Open users"', 'Opens Users setup page'],
        ['"Go to profiles"', 'Opens Profiles setup page'],
        ['"Open permission sets"', 'Opens Permission Sets page'],
        ['"Go to flows"', 'Opens Flow Builder list'],
        ['"Open apex classes"', 'Opens Apex Classes page'],
        ['"Go to custom metadata"', 'Opens Custom Metadata Types'],
        ['"Open sandboxes"', 'Opens Sandboxes page'],
        ['"Go to installed packages"', 'Opens Installed Packages'],
    ]
)

doc.add_heading('Object Manager (80+ subpages)', level=2)
add_table(
    ['Say This', 'What Happens'],
    [
        ['"Account fields"', 'Opens Account Fields & Relationships'],
        ['"Contact page layouts"', 'Opens Contact Page Layouts'],
        ['"Lead validation rules"', 'Opens Lead Validation Rules'],
        ['"Account triggers"', 'Opens Account Apex Triggers'],
        ['"Create field on account"', 'Opens Account New Field wizard'],
    ]
)

doc.add_heading('Special Pages & Keyboard Shortcut', level=2)
add_table(
    ['Say This / Press This', 'What Happens'],
    [
        ['"Go home"', 'Opens Home page'],
        ['"Open chatter"', 'Opens Chatter feed'],
        ['"Open files"', 'Opens Files page'],
        ['"Open calendar"', 'Opens Calendar'],
        ['Option + Space (Mac)', 'Toggle voice listening on/off'],
        ['Ctrl + Space (Windows)', 'Toggle voice listening on/off'],
    ]
)

doc.add_page_break()

# ── 10. Security ───────────────────────────────────────────────────────
doc.add_heading('10. Security & Permissions', level=1)

doc.add_heading('What Voice Navigator Accesses', level=2)
add_table(
    ['Access Type', 'Details', 'Security Model'],
    [
        ['Custom Objects', 'Reads object metadata via Schema.getGlobalDescribe()', 'Respects object-level permissions'],
        ['Object Fields', 'Reads field metadata', 'Respects Field-Level Security (FLS)'],
        ['Record Search', 'SOSL search by Name field', 'Respects sharing rules & record access'],
        ['Microphone', 'Browser audio input', 'User must grant permission per domain'],
    ]
)

doc.add_heading('What It Does NOT Access', level=2)
no_access = [
    'No external APIs or third-party services',
    'No API keys, tokens, or secrets',
    'No Named Credentials or Remote Site Settings',
    'No data modification — all Apex methods are read-only',
    'No DML operations (no insert, update, or delete)',
]
for item in no_access:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Permission Set Setup (Optional)', level=2)
doc.add_paragraph(
    'For orgs with strict security, you can create a Permission Set to control access:'
)
perm_steps = [
    'Go to Setup > Permission Sets > New',
    'Name: "Voice Navigator Access"',
    'Add Apex Class Access: VoiceNavigatorController',
    'Add Visualforce Page Access: VoiceRecognitionPage',
    'Assign the permission set to users who need Voice Navigator',
]
for i, s in enumerate(perm_steps, 1):
    doc.add_paragraph(f'{i}. {s}')

# ── 11. Troubleshooting ────────────────────────────────────────────────
doc.add_heading('11. Troubleshooting', level=1)

issues = [
    ('Speech recognition not supported',
     'Browser does not support Web Speech API.',
     'Switch to Google Chrome desktop browser.'),
    ('Microphone permission denied',
     'Browser microphone access was blocked.',
     'Click the lock icon in Chrome\'s address bar > Set Microphone to "Allow" > Refresh the page.'),
    ('Component not visible in App Builder',
     'Deployment may not have completed or metadata is missing.',
     'Verify deployment succeeded. Check that isExposed is true in voiceNavigator.js-meta.xml. Clear browser cache with Ctrl+Shift+R.'),
    ('Commands not recognized',
     'Speech recognition accuracy depends on microphone quality and background noise.',
     'Speak clearly at normal pace. Check the "Last Command" display. Aim for confidence above 70%. Reduce background noise.'),
    ('"No results found" for search',
     'Search term too short or record doesn\'t exist.',
     'Minimum 2-character search term required. Search only matches the Name field. Check record is accessible to your user profile.'),
    ('Custom objects not recognized',
     'Object label doesn\'t match spoken words.',
     'Say the exact object label as shown in Salesforce. The component auto-loads custom objects on initialization. Fuzzy matching helps with minor variations.'),
    ('Voice Navigator not in utility bar',
     'Component was not added to the app\'s utility items.',
     'Follow Section 5 steps. Make sure you are using the correct Lightning App. Refresh the browser after saving.'),
]

for title, cause, fix in issues:
    p = doc.add_paragraph()
    run = p.add_run(f'Issue: {title}')
    run.bold = True
    p = doc.add_paragraph()
    run = p.add_run('Cause: ')
    run.bold = True
    p.add_run(cause)
    p = doc.add_paragraph()
    run = p.add_run('Fix: ')
    run.bold = True
    p.add_run(fix)
    doc.add_paragraph()

# ── Save ────────────────────────────────────────────────────────────────
output_path = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    'Voice_Navigator_Technical_Setup_Guide.docx'
)
doc.save(output_path)
print(f'Document saved to: {output_path}')
